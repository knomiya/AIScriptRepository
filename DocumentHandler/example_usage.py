"""
DocumentHandler 使用示例
"""
from document_handler import DocumentHandler
import os

def example_excel_processing():
    """Excel处理示例"""
    print("=== Excel处理示例 ===")
    
    handler = DocumentHandler()
    
    # 示例需求
    examples = [
        ("sample_data.xlsx", "统计每列的平均值和总和"),
        ("sales_data.xlsx", "按销售额排序并筛选出前10名"),
        ("employee_data.xlsx", "分析员工年龄分布和薪资统计")
    ]
    
    for file_path, requirement in examples:
        print(f"\n处理文件: {file_path}")
        print(f"需求: {requirement}")
        
        if os.path.exists(file_path):
            result = handler.process_document(file_path, requirement)
            print(f"结果: {'成功' if result.get('success') else '失败'}")
            
            if result.get('success'):
                ai_analysis = result.get('ai_analysis', {})
                operations = ai_analysis.get('operations', [])
                print(f"识别操作: {[op.get('action') for op in operations]}")
        else:
            print(f"文件不存在，跳过处理")

def example_word_processing():
    """Word处理示例"""
    print("\n=== Word处理示例 ===")
    
    handler = DocumentHandler()
    
    # 示例需求
    examples = [
        ("report.docx", "调整文档格式，统一字体和段落样式"),
        ("contract.docx", "替换所有的甲方为新公司名称"),
        ("manual.docx", "分析文档结构并生成目录")
    ]
    
    for file_path, requirement in examples:
        print(f"\n处理文件: {file_path}")
        print(f"需求: {requirement}")
        
        if os.path.exists(file_path):
            result = handler.process_document(file_path, requirement)
            print(f"结果: {'成功' if result.get('success') else '失败'}")
            
            if result.get('success'):
                ai_analysis = result.get('ai_analysis', {})
                operations = ai_analysis.get('operations', [])
                print(f"识别操作: {[op.get('action') for op in operations]}")
        else:
            print(f"文件不存在，跳过处理")

def create_sample_files():
    """创建示例文件用于测试"""
    print("=== 创建示例文件 ===")
    
    # 创建示例Excel文件
    try:
        import pandas as pd
        
        # 示例数据
        data = {
            '姓名': ['张三', '李四', '王五', '赵六', '钱七'],
            '年龄': [25, 30, 35, 28, 32],
            '部门': ['技术部', '销售部', '技术部', '人事部', '销售部'],
            '薪资': [8000, 12000, 15000, 9000, 11000]
        }
        
        df = pd.DataFrame(data)
        df.to_excel('sample_data.xlsx', index=False)
        print("创建示例Excel文件: sample_data.xlsx")
        
    except ImportError:
        print("需要安装pandas和openpyxl来创建Excel示例文件")
    
    # 创建示例Word文件
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('示例文档', 0)
        
        doc.add_heading('第一章 概述', level=1)
        doc.add_paragraph('这是一个示例文档，用于测试DocumentHandler的功能。')
        
        doc.add_heading('第二章 详细内容', level=1)
        doc.add_paragraph('这里包含了详细的内容描述。')
        doc.add_paragraph('可以进行格式调整、内容替换等操作。')
        
        doc.save('sample_document.docx')
        print("创建示例Word文件: sample_document.docx")
        
    except ImportError:
        print("需要安装python-docx来创建Word示例文件")

if __name__ == "__main__":
    print("DocumentHandler 使用示例")
    print("=" * 40)
    
    # 创建示例文件
    create_sample_files()
    
    # 运行示例
    example_excel_processing()
    example_word_processing()
    
    print("\n" + "=" * 40)
    print("示例运行完成!")
    print("\n使用方法:")
    print("1. 安装依赖: pip install -r requirements.txt")
    print("2. 命令行使用: python document_handler.py <文件路径> '<需求描述>'")
    print("3. 代码调用: 参考本文件中的示例代码")