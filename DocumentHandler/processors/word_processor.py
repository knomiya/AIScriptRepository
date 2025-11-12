"""
Word文档处理器
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import os
import re

class WordProcessor:
    def __init__(self, file_path: str):
        """初始化Word处理器"""
        self.file_path = file_path
        self.document = None
        self.load_file()
    
    def load_file(self):
        """加载Word文件"""
        try:
            self.document = Document(self.file_path)
            print(f"成功加载Word文档，包含 {len(self.document.paragraphs)} 个段落")
            
        except Exception as e:
            raise Exception(f"加载Word文件失败: {e}")
    
    def get_file_info(self) -> Dict[str, Any]:
        """获取文件基本信息"""
        paragraphs = self.document.paragraphs
        tables = self.document.tables
        
        # 统计文字内容
        total_text = ""
        for paragraph in paragraphs:
            total_text += paragraph.text + "\n"
        
        word_count = len(total_text.split())
        char_count = len(total_text)
        
        info = {
            "file_path": self.file_path,
            "file_size": os.path.getsize(self.file_path),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "word_count": word_count,
            "character_count": char_count,
            "has_images": self._count_images(),
            "styles_used": self._get_styles_used()
        }
        
        return info
    
    def execute_operations(self, operations: List[Dict]) -> Dict[str, Any]:
        """执行处理操作"""
        results = []
        
        for operation in operations:
            try:
                result = self._execute_single_operation(operation)
                results.append({
                    "operation": operation,
                    "success": True,
                    "result": result
                })
            except Exception as e:
                results.append({
                    "operation": operation,
                    "success": False,
                    "error": str(e)
                })
        
        return {"results": results}
    
    def _execute_single_operation(self, operation: Dict) -> Any:
        """执行单个操作"""
        op_type = operation.get("type")
        
        if op_type == "format":
            return self._format_document(operation.get("parameters", {}))
        elif op_type == "replace":
            return self._replace_content(operation.get("parameters", {}))
        elif op_type == "analyze":
            return self._analyze_content()
        else:
            return f"未知操作类型: {op_type}"
    
    def _format_document(self, parameters: Dict) -> Dict[str, Any]:
        """格式化文档"""
        format_results = {
            "formatted_paragraphs": 0,
            "changes_made": []
        }
        
        # 示例格式化操作
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():  # 非空段落
                # 设置基本格式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 设置字体
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                
                format_results["formatted_paragraphs"] += 1
        
        format_results["changes_made"].append("设置段落左对齐")
        format_results["changes_made"].append("设置字体为宋体12号")
        
        return format_results
    
    def _replace_content(self, parameters: Dict) -> Dict[str, Any]:
        """替换内容"""
        # 这里可以根据参数实现具体的替换逻辑
        replace_results = {
            "replacements_made": 0,
            "patterns_found": []
        }
        
        # 示例：查找常见的格式问题
        for paragraph in self.document.paragraphs:
            text = paragraph.text
            
            # 查找多余的空格
            if "  " in text:  # 两个或更多空格
                replace_results["patterns_found"].append("发现多余空格")
            
            # 查找中英文混排问题
            if re.search(r'[a-zA-Z][^\s][中文]', text):
                replace_results["patterns_found"].append("发现中英文混排问题")
        
        return replace_results
    
    def _analyze_content(self) -> Dict[str, Any]:
        """分析文档内容"""
        analysis = {
            "document_structure": self._analyze_structure(),
            "content_statistics": self._get_content_statistics(),
            "formatting_analysis": self._analyze_formatting()
        }
        
        return analysis
    
    def _analyze_structure(self) -> Dict[str, Any]:
        """分析文档结构"""
        structure = {
            "headings": [],
            "paragraphs_by_level": {},
            "tables": []
        }
        
        for paragraph in self.document.paragraphs:
            style_name = paragraph.style.name
            if 'Heading' in style_name:
                structure["headings"].append({
                    "level": style_name,
                    "text": paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
                })
        
        # 分析表格
        for i, table in enumerate(self.document.tables):
            structure["tables"].append({
                "table_index": i,
                "rows": len(table.rows),
                "columns": len(table.columns)
            })
        
        return structure
    
    def _get_content_statistics(self) -> Dict[str, Any]:
        """获取内容统计"""
        all_text = ""
        for paragraph in self.document.paragraphs:
            all_text += paragraph.text + " "
        
        # 简单的内容分析
        sentences = all_text.split('。')
        words = all_text.split()
        
        return {
            "total_characters": len(all_text),
            "total_words": len(words),
            "total_sentences": len([s for s in sentences if s.strip()]),
            "average_sentence_length": len(all_text) / max(len(sentences), 1),
            "paragraphs_with_content": len([p for p in self.document.paragraphs if p.text.strip()])
        }
    
    def _analyze_formatting(self) -> Dict[str, Any]:
        """分析格式"""
        fonts_used = set()
        font_sizes = set()
        alignments = set()
        
        for paragraph in self.document.paragraphs:
            alignments.add(str(paragraph.alignment))
            
            for run in paragraph.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    font_sizes.add(str(run.font.size))
        
        return {
            "fonts_used": list(fonts_used),
            "font_sizes": list(font_sizes),
            "alignments": list(alignments)
        }
    
    def _count_images(self) -> int:
        """统计图片数量"""
        # 简化实现，实际可能需要更复杂的逻辑
        return 0
    
    def _get_styles_used(self) -> List[str]:
        """获取使用的样式"""
        styles = set()
        for paragraph in self.document.paragraphs:
            styles.add(paragraph.style.name)
        return list(styles)
    
    def save_processed_file(self, output_path: str):
        """保存处理后的文件"""
        try:
            self.document.save(output_path)
            print(f"处理后的文件已保存到: {output_path}")
            return True
            
        except Exception as e:
            print(f"保存文件失败: {e}")
            return False